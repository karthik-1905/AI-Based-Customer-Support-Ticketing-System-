"""
Resume Parser - Extract raw text from PDF and DOCX resume files.

Uses:
  - pdfplumber for PDF files (preserves layout better)
  - PyMuPDF (fitz) as fallback for PDFs
  - python-docx for .docx files
"""
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract all text from a PDF file using pdfplumber.
    Falls back to PyMuPDF (fitz) if pdfplumber fails.

    Args:
        file_path: Absolute or relative path to the PDF file.

    Returns:
        Concatenated text from all pages, or empty string on failure.
    """
    text_parts: list[str] = []

    # --- Primary: pdfplumber ---
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        if text_parts:
            return "\n".join(text_parts)
    except Exception as exc:
        logger.warning(f"pdfplumber failed for {file_path}: {exc}")

    # --- Fallback: PyMuPDF ---
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except Exception as exc:
        logger.error(f"PyMuPDF also failed for {file_path}: {exc}")

    return ""


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a DOCX file using python-docx.

    Args:
        file_path: Absolute or relative path to the DOCX file.

    Returns:
        Concatenated paragraph text, or empty string on failure.
    """
    try:
        from docx import Document  # python-docx

        doc = Document(file_path)
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        return "\n".join(paragraphs)
    except Exception as exc:
        logger.error(f"python-docx failed for {file_path}: {exc}")
        return ""


def extract_text(file_path: str) -> str:
    """
    Main entry point - dispatch to the correct extractor based on file extension.

    Args:
        file_path: Path to the resume file (.pdf or .docx).

    Returns:
        Extracted raw text. Returns empty string on failure.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        text = extract_text_from_docx(file_path)
    else:
        logger.warning(f"Unsupported file type: {ext}")
        return ""

    # Normalise whitespace - collapse multiple blank lines
    import re
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)

    return text.strip()
